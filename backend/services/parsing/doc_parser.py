from docx import Document


def extract_full_text(file_path: str) -> str:
    """Extract all text from a Word document as a single string."""
    doc = Document(file_path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()


def extract_text_pages(file_path: str, print_paragraphs: bool = False) -> list[dict]:
    """
    Extract text from a Word document paragraph by paragraph.
    
    Returns a list of dicts: [{"paragraph": 1, "text": "..."}, ...]
    Optionally prints each paragraph with clear formatting.
    """
    doc = Document(file_path)
    paragraphs = []
    
    for i, para in enumerate(doc.paragraphs, start=1):
        text = para.text or ""
        para_data = {"paragraph": i, "text": text.strip()}
        paragraphs.append(para_data)
        
    return paragraphs


# Simple alias
parse_docx = extract_full_text